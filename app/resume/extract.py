"""简历文本抽取。

把上传的简历文件（PDF / Word / Markdown / 纯文本）抽取为纯文本，
供后续 ``parser.py`` 结构化解析使用。

本模块只做抽取，不调用模型、不下载任何资源；缺失可选依赖时抛出
明确的 :class:`ExtractionError`，由上层提示安装对应依赖。
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import Optional

# 支持的文件扩展名 -> 来源类型（写入 Resume.source_type）
SUPPORTED_EXTENSIONS = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".md": "md",
    ".markdown": "md",
    ".txt": "txt",
}

# 上传大小上限（字节）
MAX_FILE_BYTES = 10 * 1024 * 1024  # 10MB


class ExtractionError(RuntimeError):
    """文本抽取失败（不支持的类型 / 依赖缺失 / 文件损坏等）。"""


def detect_source_type(filename: str) -> str:
    """根据文件名后缀返回来源类型，不支持时抛出 ExtractionError。"""
    suffix = Path(filename).suffix.lower()
    source_type = SUPPORTED_EXTENSIONS.get(suffix)
    if source_type is None:
        supported = "、".join(sorted(SUPPORTED_EXTENSIONS))
        raise ExtractionError(f"不支持的文件类型「{suffix}」，仅支持：{supported}")
    return source_type


def _extract_pdf(path: Path) -> str:
    """从 PDF 抽取纯文本。

    优先使用 ``pypdf``（纯 Python，依赖轻）；若不可用则回退到
    ``pdfplumber``。两者都缺失时抛出 ExtractionError。
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        return _extract_pdf_plumber(path)

    try:
        reader = PdfReader(str(path))
        parts = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text.strip())
        text = "\n\n".join(parts)
    except Exception as exc:  # noqa: BLE001 - 文件损坏等异常统一包装
        raise ExtractionError(f"PDF 解析失败：{exc}") from exc

    if not text.strip():
        raise ExtractionError("PDF 未抽取到任何文本，可能是扫描件（图片型 PDF）")
    return text


def _extract_pdf_plumber(path: Path) -> str:
    """使用 pdfplumber 兜底抽取 PDF 文本。"""
    try:
        import pdfplumber  # noqa: F401
    except ImportError:
        raise ExtractionError(
            "缺少 PDF 解析依赖，请安装 pypdf（推荐）或 pdfplumber"
        ) from None

    try:
        import pdfplumber as plumber

        parts = []
        with plumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                if text.strip():
                    parts.append(text.strip())
        text = "\n\n".join(parts)
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"PDF 解析失败：{exc}") from exc

    if not text.strip():
        raise ExtractionError("PDF 未抽取到任何文本，可能是扫描件（图片型 PDF）")
    return text


def _extract_docx(path: Path) -> str:
    """从 Word (.docx) 抽取纯文本。"""
    try:
        from docx import Document
    except ImportError:
        raise ExtractionError("缺少 Word 解析依赖，请安装 python-docx") from None

    try:
        doc = Document(str(path))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # 表格内容一并抽取（简历中的技能/项目常以表格呈现）
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"Word 解析失败：{exc}") from exc

    if not text.strip():
        raise ExtractionError("Word 文档未抽取到任何文本")
    return text


def _extract_markdown(path: Path) -> str:
    """读取 Markdown / 纯文本（直接返回原文）。"""
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="gbk", errors="replace")
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"文本读取失败：{exc}") from exc
    return text


def extract_text_from_bytes(content: bytes, source_type: str) -> str:
    """从内存字节中抽取纯文本（供上传接口直接使用，无需先落盘）。

    文本类（md/txt）直接解码；PDF/Word 先写入临时文件再解析
    （pypdf/python-docx 需要真实文件对象或文件路径）。

    Args:
        content: 文件原始字节。
        source_type: 来源类型（pdf/docx/md/txt）。

    Returns:
        抽取出的纯文本。

    Raises:
        ExtractionError: 依赖缺失或解析失败。
    """
    if source_type in {"md", "txt"}:
        text = _decode_text_bytes(content)
        if not text.strip():
            raise ExtractionError("文本文件内容为空")
        return text

    if source_type not in {"pdf", "docx"}:
        raise ExtractionError(f"未知来源类型：{source_type}")

    # PDF/Word 需要真实文件：写入临时文件，解析后清理
    import tempfile

    suffix = ".pdf" if source_type == "pdf" else ".docx"
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        if source_type == "pdf":
            return _extract_pdf(Path(tmp_path))
        return _extract_docx(Path(tmp_path))
    finally:
        try:
            os.remove(tmp_path)
        except OSError:
            pass


def _decode_text_bytes(content: bytes) -> str:
    """尝试按 UTF-8 解码，失败回退 GBK。"""
    for encoding in ("utf-8", "gbk"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def extract_text(path: str, source_type: Optional[str] = None) -> str:
    """抽取指定文件的纯文本。

    Args:
        path: 文件路径。
        source_type: 来源类型（pdf/docx/md/txt）；为空时按后缀自动识别。

    Returns:
        抽取出的纯文本。

    Raises:
        ExtractionError: 类型不支持、依赖缺失或解析失败。
    """
    file_path = Path(path)
    if not file_path.exists():
        raise ExtractionError(f"文件不存在：{path}")

    if source_type is None:
        source_type = detect_source_type(file_path.name)

    if source_type == "pdf":
        return _extract_pdf(file_path)
    if source_type == "docx":
        return _extract_docx(file_path)
    if source_type in {"md", "txt"}:
        return _extract_markdown(file_path)

    raise ExtractionError(f"未知来源类型：{source_type}")
